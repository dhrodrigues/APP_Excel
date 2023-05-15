import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tqdm import tqdm
import tkinter as tk
from tkinter import filedialog
import time

start_time = time.time()

# abre arquivos office
root = tk.Tk()
root.withdraw()
file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
workbook = openpyxl.load_workbook(file_path)
worksheet = workbook['Planilha1']

root = tk.Tk()
root.withdraw()
file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
doc = Document(file_path)

# iterar sobre as tabelas no documento
for table in tqdm(doc.tables):

    # iterar sobre os dados da planilha e preencher a tabela do Word
    for row in worksheet.iter_rows(min_row=2, max_col=2, values_only=True):
        pos_no = row[0]
        descricao = row[1]
        
        # procurar por correspondências na tabela do Word
        for i, row_table in enumerate(table.rows):
            if row_table.cells[0].text.strip() == str(pos_no):
                
                cell = table.cell(i, 1)                
                empty_paragraph_before = cell.add_paragraph()
                new_paragraph = cell.add_paragraph(descricao)
                empty_paragraph_after = cell.add_paragraph()              


                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # definir a fonte do parágrafo
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(7)

                # adicionar um tabulador para alinhar à esquerda
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'left')
                new_paragraph._element.append(tab)

                # definir o recuo para direita e esquerda
                new_paragraph_format = new_paragraph.paragraph_format
                new_paragraph_format.left_indent = Inches(0.05)
                new_paragraph_format.right_indent = Inches(0.05)

                original_cell = row_table.cells[0]._element
                break

# salvar as alterações no arquivo do Word
root = tk.Tk()
root.withdraw()
save_file_path = filedialog.asksaveasfilename(filetypes=[("Word files", "*.docx")])
doc.save(save_file_path)

# mostrar o tempo de execução
end_time = time.time()
total_time = end_time - start_time
print(f"Tempo total: {total_time} segundos")
